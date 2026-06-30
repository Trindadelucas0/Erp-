"""Gera documento Word com cronograma de implantação ERP (estado real do projeto)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "docs"
OUTPUT = OUTPUT_DIR / "cronograma_implantacao_erp.docx"
OUTPUT_FALLBACK = OUTPUT_DIR / "cronograma_implantacao_erp_v2.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = cell_text
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def main() -> None:
    doc = Document()

    title = doc.add_heading("Cronograma de Implantação ERP", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Fases, entregas e homologação")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    add_paragraph(
        doc,
        "Documento atualizado com base no estado real do sistema em 30/06/2026. "
        "Reorganiza as datas indicadas pelo cliente (29/06 a 10/08) em fases com "
        "dependência técnica, janelas de homologação e ritmo sustentável de entrega.",
    )

    add_heading(doc, "Premissas gerais", 2)
    for item in [
        "Cada fase inclui: desenvolvimento → entrega para testes → homologação (5 dias úteis) → correções.",
        "Atrasos na homologação ou mudanças de escopo deslocam as fases seguintes.",
        "O escopo de cada fase será detalhado em documento de aceite antes do início.",
        "Importação em massa (XLS, planilhas legadas) será tratada após as telas de cadastro estarem homologadas, salvo acordo específico.",
        "Um ponto focal do cliente para decisões de negócio evita retrabalho.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Módulos já entregues e operacionais", 2)
    add_paragraph(
        doc,
        "Estes módulos já possuem backend, frontend e persistência implementados. "
        "Não fazem parte do cronograma futuro — apenas homologação ou ajustes pontuais, quando indicado.",
    )
    add_table(
        doc,
        ["Módulo", "O que está pronto"],
        [
            [
                "Infraestrutura",
                "Autenticação JWT, RBAC, auditoria, validações CPF/CNPJ/IE, estrutura Fastify + Next.js",
            ],
            [
                "Empresas",
                "CRUD multiempresa, vínculo usuário-empresa, tela /cadastros",
            ],
            [
                "Clientes",
                "CRUD PF/PJ, busca automática CNPJ (BrasilAPI), múltiplos contatos e endereços, "
                "fluxo de aprovação PJ, assinatura digital (ZapSign), tela /clientes/aprovacao",
            ],
            [
                "Fornecedores",
                "CRUD completo com abas, duplicidade inteligente por documento, vínculos entre fornecedores, "
                "seleção de planos financeiros e CFOP (catálogo), pares plano+CFOP padrão, "
                "flags operacionais (revenda, consumo, vínculo manual de produtos na entrada etc.)",
            ],
            [
                "Transportadoras",
                "CRUD completo, duplicidade por documento, campos ANTT e aceita NF-e modelo 55",
            ],
            [
                "Administração",
                "Usuários, papéis, permissões por módulo, páginas vinculáveis por usuário, atalhos de teclado",
            ],
            [
                "Assinatura digital",
                "Integração ZapSign: configuração, envio, listagem, webhook e página pública de assinatura",
            ],
        ],
    )

    add_heading(doc, "Situação atual do projeto", 2)
    add_table(
        doc,
        ["Área", "Situação"],
        [
            ["Infraestrutura, Empresas e Administração", "Entregue e operacional"],
            ["Clientes", "Entregue e operacional (PF/PJ, aprovação, ZapSign)"],
            [
                "Fornecedores e Transportadoras",
                "Entregue; em homologação — ajustes finos e pendências do cliente",
            ],
            [
                "Planos Financeiros e CFOP",
                "Parcial — tabelas no banco + catálogo somente leitura usado no fornecedor; "
                "telas de gestão (CRUD, árvore, layout definitivo) pendentes",
            ],
            [
                "Produtos, NCM, Marcas, Norma de Palete",
                "Não iniciado — documentado em blocos-de-entrega; sem models, rotas ou telas",
            ],
            [
                "Financeiro (contas a pagar/receber)",
                "Não iniciado — permissões pré-configuradas; sem lançamentos ou telas",
            ],
            [
                "Compras, Entrada de NF, Estoque",
                "Não iniciado — permissões pré-configuradas; sem models operacionais",
            ],
        ],
    )

    add_heading(doc, "Fase 0 — Homologação de Fornecedores e Transportadoras", 2)
    add_paragraph(doc, "Período: até 11/07/2026", bold=True)
    add_paragraph(
        doc,
        "Os módulos já foram desenvolvidos. Esta fase é de validação com o cliente, "
        "não de construção do zero.",
    )
    add_table(
        doc,
        ["Entrega", "Conteúdo"],
        [
            [
                "Fornecedores (homologação)",
                "Validar CRUD, abas, duplicidade, vínculos entre fornecedores, "
                "planos/CFOP no formulário, pares plano+CFOP padrão e flags operacionais",
            ],
            [
                "Transportadoras (homologação)",
                "Validar CRUD, duplicidade por documento, campos ANTT e aceita NF-e 55",
            ],
            [
                "Ajustes finais",
                "Correções de bugs, refinamentos de UX e regras apontados pelo cliente "
                "(ex.: padronização de máscaras e caixa alta em campos de texto)",
            ],
        ],
    )
    add_paragraph(
        doc,
        "Objetivo: encerrar homologação dos cadastros de pessoas (cliente já entregue; "
        "fornecedor e transportadora nesta fase) antes de iniciar telas fiscais e produtos.",
    )

    add_heading(doc, "Fase 1 — Base fiscal (Planos Financeiros e CFOP)", 2)
    add_paragraph(doc, "Período: 14/07 a 25/07/2026 | Homologação: 28/07 a 01/08", bold=True)
    add_table(
        doc,
        ["Módulo", "Entrega"],
        [
            [
                "Planos Financeiros",
                "Telas de gestão: árvore Receitas/Despesas/Resultado, numeração hierárquica, "
                "CRUD completo, integração com fornecedor (substitui catálogo estático atual)",
            ],
            [
                "CFOP",
                "Telas de gestão: cadastro completo (layout ERP), tipos 01–06, "
                "vínculo com fornecedor (substitui catálogo estático atual)",
            ],
        ],
    )
    add_paragraph(
        doc,
        "Hoje planos e CFOP existem apenas como tabelas e endpoints de catálogo (leitura) "
        "para seleção no fornecedor. Esta fase entrega o cadastro mestre completo.",
    )

    add_heading(doc, "Fase 2 — Cadastros de Produtos e auxiliares", 2)
    add_paragraph(doc, "Período: 04/08 a 22/08/2026 | Homologação: 25/08 a 29/08", bold=True)
    add_table(
        doc,
        ["Módulo", "Entrega"],
        [
            ["NCM", "Cadastro e consulta (base para produto e NF)"],
            ["Marcas", "CRUD"],
            ["Norma de Palete", "CRUD e vínculo com produto"],
            [
                "Produtos",
                "Cadastro completo (dados fiscais, marca, palete, fornecedor, estoque mínimo etc.)",
            ],
        ],
    )
    add_paragraph(
        doc,
        "NCM, Marcas e Norma de Palete entram antes do cadastro de Produto para evitar refatoração. "
        "Depende de: Fase 1 (CFOP) e cadastros de pessoas homologados.",
    )

    add_heading(doc, "Fase 3 — Financeiro (Contas a Pagar e a Receber)", 2)
    add_paragraph(doc, "Período: 01/09 a 19/09/2026 | Homologação: 22/09 a 26/09", bold=True)
    add_table(
        doc,
        ["Módulo", "Entrega"],
        [
            [
                "Contas a pagar",
                "Lançamento, vencimento, plano financeiro, fornecedor, anexos (conforme regra do plano)",
            ],
            ["Contas a receber", "Lançamento, cliente, recebimento"],
        ],
    )
    add_paragraph(
        doc,
        "Depende de: Planos Financeiros com CRUD (Fase 1), Fornecedores e Clientes (já entregues).",
    )

    add_heading(doc, "Fase 4 — Pedido de Compras", 2)
    add_paragraph(doc, "Período: 29/09 a 17/10/2026 | Homologação: 20/10 a 24/10", bold=True)
    add_table(
        doc,
        ["Entrega", "Conteúdo"],
        [
            [
                "Pedido de compra",
                "Cabeçalho, itens, fornecedor, produtos, status (rascunho/aprovado/cancelado)",
            ],
            ["Integração", "Vínculo com fornecedores e produtos homologados"],
        ],
    )

    add_heading(doc, "Fase 5 — Entrada de Notas e Controle de Estoque", 2)
    add_paragraph(doc, "Período: 27/10 a 21/11/2026 | Homologação: 24/11 a 28/11", bold=True)
    add_table(
        doc,
        ["Módulo", "Entrega"],
        [
            [
                "Entrada de NF",
                "Importação/leitura XML, vínculo fornecedor/produto, CFOP, conferência",
            ],
            ["Estoque", "Movimentação por entrada, saldo por produto, histórico básico"],
        ],
    )
    add_paragraph(
        doc,
        "Módulo mais complexo: exige Produtos, CFOP, Fornecedores e Pedido de Compras estáveis.",
    )

    add_heading(doc, "Fase 6 — Consolidação e estabilização", 2)
    add_paragraph(doc, "Período: 01/12 a 19/12/2026", bold=True)
    for item in [
        "Correção dos principais bugs das fases anteriores",
        "Ajustes de UX e regras de negócio pendentes",
        "Testes integrados (compra → entrada → estoque → financeiro)",
        "Documentação resumida de uso",
        "Go-live operacional do conjunto de módulos entregues",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Visão resumida do cronograma", 2)
    add_table(
        doc,
        ["Período", "Entrega principal"],
        [
            ["Concluído", "Infraestrutura, Empresas, Clientes, Admin, ZapSign"],
            ["Jul/2026", "Homologação Fornecedor e Transportadora + início Planos/CFOP"],
            ["Jul–Ago/2026", "Planos Financeiros + CFOP (telas de gestão)"],
            ["Ago/2026", "Produtos + NCM + Marcas + Norma de Palete"],
            ["Set/2026", "Financeiro (pagar/receber)"],
            ["Out/2026", "Pedido de Compras"],
            ["Nov/2026", "Entrada de NF + Estoque"],
            ["Dez/2026", "Consolidação e go-live"],
        ],
    )

    add_heading(doc, "Responsabilidades do cliente em cada fase", 2)
    for item in [
        "Validação em até 5 dias úteis após cada entrega (ou registro formal de pendências).",
        "Especificações fechadas antes do início da fase (prints, campos obrigatórios, fluxos).",
        "Dados de teste (CNPJs, produtos exemplo, XML de NF reais anonimizados).",
        "Um ponto focal para decisões de negócio.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Sobre a meta de finalização em agosto", 2)
    add_paragraph(
        doc,
        "Em agosto é viável entregar e homologar: homologação de Fornecedores e Transportadoras, "
        "Planos Financeiros e CFOP (telas de gestão), e início ou entrega parcial de Produtos "
        "(conforme complexidade dos campos confirmados com o cliente).",
    )
    add_paragraph(
        doc,
        "Clientes, infraestrutura e administração já estão operacionais e não consomem cronograma futuro. "
        "Compras, entrada de NF, estoque e financeiro completo formam um bloco integrado que, "
        "com qualidade e homologação adequadas, se posiciona entre setembro e novembro, "
        "com consolidação em dezembro.",
    )

    add_heading(doc, "Alternativa — Entrega por ondas (MVP)", 2)
    add_paragraph(
        doc,
        "Caso a prioridade seja operação mais cedo, é possível discutir entrega por ondas:",
    )
    add_table(
        doc,
        ["Onda", "Prazo", "O que entra", "O que fica para depois"],
        [
            [
                "Já entregue",
                "—",
                "Infraestrutura, Empresas, Clientes, Fornecedor, Transportadora, Admin, ZapSign",
                "—",
            ],
            [
                "Onda 1 — Cadastros fiscais e produtos",
                "até 29/08/2026",
                "Planos Financeiros, CFOP, Produtos (MVP), NCM, Marcas, Norma de Palete",
                "Importação XLS, relatórios avançados",
            ],
            [
                "Onda 2 — Compras/Estoque",
                "até 24/10/2026",
                "Pedido compra, Entrada NF, Estoque básico",
                "Regras complexas, integrações extras",
            ],
            [
                "Onda 3 — Financeiro",
                "até 21/11/2026",
                "Pagar/receber + bugs integrados",
                "Automações, dashboards",
            ],
        ],
    )

    add_heading(doc, "Comparativo com datas originais do cliente", 2)
    add_table(
        doc,
        ["Data original", "Observação"],
        [
            [
                "29/06 — Fornecedor e Transportadoras",
                "Módulos já desenvolvidos; período atual é homologação e ajustes (Fase 0 até 11/07)",
            ],
            [
                "06/07 — Produto + NCM + CFOP + Marcas + Palete",
                "CFOP ainda sem tela de gestão; produtos não iniciados; "
                "proposta: Fase 1 (jul) + Fase 2 (ago)",
            ],
            [
                "13/07 — Financeiro",
                "Depende de Planos Financeiros com CRUD; proposta: setembro",
            ],
            [
                "20/07 — Pedido de Compras",
                "Depende de produtos; proposta: outubro",
            ],
            [
                "27/07 — NF + Estoque",
                "Módulo complexo, sem código ainda; proposta: novembro",
            ],
            [
                "10/08 — Consolidação",
                "Insuficiente para homologação real de todos os módulos; proposta: dezembro",
            ],
        ],
    )

    doc.add_paragraph()
    closing = doc.add_paragraph(
        "Este cronograma reflete o que já foi entregue (clientes, fornecedores, transportadoras, "
        "infraestrutura) e organiza o trabalho restante com homologação adequada. Cadastros e base "
        "fiscal podem ser consolidados em agosto; módulos operacionais (compra, nota, estoque, "
        "financeiro) seguem em sequência até dezembro."
    )
    for run in closing.runs:
        run.font.size = Pt(11)
        run.italic = True

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    destino = OUTPUT
    try:
        doc.save(destino)
    except PermissionError:
        destino = OUTPUT_FALLBACK
        doc.save(destino)
    print(f"Documento gerado: {destino}")


if __name__ == "__main__":
    main()
